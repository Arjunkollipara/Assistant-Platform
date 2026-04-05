"""
File: backend/ingestion/parsers/docx_parser.py
Purpose: DOCX parser that extracts paragraph text from Word documents.
"""

import io

from docx import Document

from backend.ingestion.parsers.base import ParsedDocument


def parse_docx(file_bytes: bytes) -> ParsedDocument:
    """
    Parse DOCX bytes into normalized plain text.
    """
    document = Document(io.BytesIO(file_bytes))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
    non_empty_paragraphs = [paragraph for paragraph in paragraphs if paragraph]
    text = "\n".join(non_empty_paragraphs)

    return ParsedDocument(
        document_type="docx",
        text=text,
        metadata={
            "paragraph_count": len(non_empty_paragraphs),
            "char_count_raw": len(text),
        },
    )

